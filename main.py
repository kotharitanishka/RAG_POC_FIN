"""
RAG (Retrieval-Augmented Generation) Proof of Concept
Uses Redis Vector Search with LangChain and Anthropic Claude

Optimized with lazy imports for faster startup.
"""

import os
import asyncio
import uuid
import warnings
import getpass
from dotenv import load_dotenv

load_dotenv()
api_key = os.getenv("CLAUDE_API_KEY")
gemini_api_key = os.getenv("GEMINI_API_KEY")
redis_host=os.getenv("REDIS_HOST")
redis_port=os.getenv("REDIS_PORT")

# Only lightweight imports at startup
warnings.filterwarnings("ignore")
os.environ["TOKENIZERS_PARALLELISM"] = "false"

# =============================================================================
# Configuration
# =============================================================================

REDIS_HOST = redis_host
REDIS_PORT = redis_port
REDIS_PASSWORD = ""
REDIS_URL = f"redis://:{REDIS_PASSWORD}@{REDIS_HOST}:{REDIS_PORT}"

CHAT_MODEL = "claude-sonnet-4-5-20250929"
INDEX_NAME = "redisvl"
DOC_PATH = "resources/Questionnaire_for_Improving_Liquidity_in_Bond_Market.pdf"

SYSTEM_PROMPT = """
You are a high-precision Multilingual AI Assistant. Your knowledge base consists of FAQS ,PDF documents and Transcribed Audio files in English and multiple Indian languages.

OPERATIONAL RULES:
0. RESPONSE : Keep the responses short crisp to the point .
1. GROUNDING: Answer ONLY using provided context. If information is missing, state: "I could not find this information in the provided documents."
2. LANGUAGE MATCHING: Respond in the specific language used by the user in their query; default would be english.
3. CROSS-LINGUAL BRIDGE & VERIFICATION: 
   - You must translate facts accurately from the context language into the user's query language.
   - For every major fact provided, follow it with the original text from the source in its original language, placed in brackets. 
   - Format: [Original: "Context text in original script"]
4. ACCURACY: Never invent policy details, dates, or numbers. Do not use external knowledge.
5. TERMINOLOGY: Preserve technical or financial terms. Match the user's mixing style (e.g., Hinglish/Marathlish) for clarity if they use it.
6. ENTITY STRICTNESS: Ensure the answer corresponds exactly to the entity (person/company/policy) asked in the query. Do not offer information about similar entities unless relevant to the comparison.
"""

# Index schema for Redis Vector Search
# SCHEMA = {
#     "index": {
#         "name": INDEX_NAME,
#         "prefix": "chunk",
#         "storage_type": "hash"
#     },
#     "fields": [
#         {"name": "chunk_id", "type": "tag", "attrs": {"sortable": True}},
#         {"name": "document_id", "type": "tag", "attrs": {"sortable": True}},

#         {"name": "content", "type": "text"},
#         {
#             "name": "text_embedding",
#             "type": "vector",
#             "attrs": {
#                 "dims" : 384,
#                 # "dims": 768,
#                 "distance_metric": "cosine",
#                 "algorithm": "hnsw",
#                 "datatype": "float32"
#             }
#         }
#     ]
# }
# Updated Index schema for Redis Vector Search with Metadata
SCHEMA = {
    "index": {
        "name": INDEX_NAME,
        "prefix": "chunk",
        "storage_type": "hash"
    },
    "fields": [
        {"name": "chunk_id", "type": "tag", "attrs": {"sortable": True}},
        {"name": "document_id", "type": "tag", "attrs": {"sortable": True}},
        {"name": "content", "type": "text"},
        
        # NEW METADATA FIELDS:
        {"name": "topic", "type": "tag", "attrs": {"sortable": True}},           # For filtering by topic
        {"name": "keywords", "type": "text"},                                     # For text search
        {"name": "question_type", "type": "tag", "attrs": {"sortable": True}},   # For filtering by type
        {"name": "question", "type": "text"},                                     # For display
        {"name": "source_file", "type": "tag", "attrs": {"sortable": True}},     # For tracking
        
        {
            "name": "text_embedding",
            "type": "vector",
            "attrs": {
                "dims": 384,
                "distance_metric": "cosine",
                "algorithm": "hnsw",
                "datatype": "float32"
            }
        }
    ]
}
# Global references (set during initialization)
_vectorizer = None
_async_index = None


# =============================================================================
# Lazy Import Helpers
# =============================================================================

def _import_redis():
    from redis import Redis
    return Redis


def _import_pdf_loader():
    print("Loading PDF processing libraries...")
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_community.document_loaders import PyPDFLoader
    return RecursiveCharacterTextSplitter, PyPDFLoader

def _import_split_text():
    print("Loading text splitting processing libraries...")
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    return RecursiveCharacterTextSplitter


def _import_vectorizer():
    print("Loading embedding model (this may take a minute on first run)...")
    from redisvl.utils.vectorize import HFTextVectorizer
    from redisvl.extensions.cache.embeddings import EmbeddingsCache
    return HFTextVectorizer, EmbeddingsCache


def _import_index():
    from redisvl.index import SearchIndex, AsyncSearchIndex
    from redisvl.query import VectorQuery
    from redisvl.redis.utils import array_to_buffer
    return SearchIndex, AsyncSearchIndex, VectorQuery, array_to_buffer


# =============================================================================
# Redis Connection
# =============================================================================

def get_redis_connection():
    """Establish and test Redis connection."""
    Redis = _import_redis()
    redis_connection = Redis.from_url(
        REDIS_URL,
        socket_timeout=2,
        socket_connect_timeout=2,
    )
    try:
        redis_connection.ping()
        print("✓ Redis connected successfully")
        return redis_connection
    except Exception as e:
        print(f"✗ Redis connection failed: {e}")
        raise


# =============================================================================
# Document Processing
# =============================================================================

def load_and_split_document(doc_path: str, chunk_size: int = 600, chunk_overlap: int = 150):
    """Load PDF and split into chunks."""
    RecursiveCharacterTextSplitter, PyPDFLoader = _import_pdf_loader()
    
    loader = PyPDFLoader(doc_path)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    chunks = loader.load_and_split(text_splitter)
    print(f"✓ Created {len(chunks)} chunks from: {doc_path}")
    return chunks

def load_and_split_docx(doc_path: str, chunk_size: int = 400, chunk_overlap: int = 100):
    """Load DOCX and split into chunks."""
    RecursiveCharacterTextSplitter = _import_split_text()
    # Import docx library
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    # Read docx file
    doc = Document(doc_path)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    # Split into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    chunks = text_splitter.create_documents([text])
    print(f"✓ Created {len(chunks)} chunks from: {doc_path}")
    return chunks
    

def load_and_split_docx_new(doc_path: str):
    """
    Load DOCX and split into FAQ chunks with metadata.
    
    Metadata includes:
    - topic: Section/category name
    - keywords: Auto-extracted important terms
    - question_type: Type of question (how-to, definition, etc.)
    - question: Original question text
    - source_file: Document filename
    - chunk_id: Unique identifier
    """
    
    
    
    from docx import Document as DocxDocument
    from langchain_core.documents import Document as LangchainDocument
    
    doc = DocxDocument(doc_path)

    faqs = []
    current_topic = "General"  # Default topic
    current_question = None
    current_answer = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect topic headings (Heading styles or ALL CAPS text)
        if para.style.name.startswith('Heading2') or text.isupper() or para.style.name.startswith('Heading 1'):
            current_topic = text
            print(f"  → Topic detected: {current_topic}")
            continue

        # Question detection
        if text.endswith("?"):
            # Save previous FAQ
            if current_question and current_answer:
                answer_text = ' '.join(current_answer).strip()
                
                faqs.append(
                    LangchainDocument(
                        page_content=f"{current_question}\n{answer_text}",
                        metadata={
                            "topic": current_topic,
                            "keywords": extract_keywords(current_question, answer_text),
                            "question_type": classify_question_type(current_question),
                            "question": current_question,
                            "source_file": doc_path.split('/')[-1].split('\\')[-1],
                            "chunk_id": f"faq_{len(faqs):04d}",
                        }
                    )
                )

            # Start new FAQ
            current_question = text
            current_answer = []

        else:
            current_answer.append(text)

    # Save last FAQ
    if current_question and current_answer:
        answer_text = ' '.join(current_answer).strip()
        faqs.append(
            LangchainDocument(
                page_content=f"{current_question}\n{answer_text}",
                metadata={
                    "topic": current_topic,
                    "keywords": extract_keywords(current_question, answer_text),
                    "question_type": classify_question_type(current_question),
                    "question": current_question,
                    "source_file": doc_path.split('/')[-1].split('\\')[-1],
                    "chunk_id": f"faq_{len(faqs):04d}",
                }
            )
        )

    # Print summary
    topics = set(faq.metadata['topic'] for faq in faqs)
    print(f"\n✓ Created {len(faqs)} FAQ chunks from: {doc_path}")
    print(f"✓ Topics found: {', '.join(topics)}")
    print(f"✓ Question types: {', '.join(set(faq.metadata['question_type'] for faq in faqs))}")
    
    return faqs


def extract_keywords(question: str, answer: str) -> list:
    import re
    """Extract important keywords from question and answer."""
    # Combine question and answer
    text = f"{question} {answer}".lower()
    
    # Common stopwords to exclude
    stopwords = {
        'is', 'the', 'a', 'an', 'in', 'to', 'of', 'for', 'on', 'with',
        'what', 'how', 'can', 'will', 'i', 'you', 'be', 'do', 'does',
        'are', 'at', 'by', 'from', 'or', 'as', 'this', 'that', 'it',
        'have', 'has', 'not', 'but', 'if', 'when', 'which', 'we', 'they'
    }
    
    # Extract words (2+ chars, keep acronyms like API, KRA, BSDA)
    words = re.findall(r'\b[A-Za-z0-9]{2,}\b', text)
    
    # Filter stopwords and deduplicate while preserving order
    seen = set()
    keywords = []
    for word in words:
        word_lower = word.lower()
        if word_lower not in stopwords and word_lower not in seen:
            keywords.append(word)
            seen.add(word_lower)
    
    # Return top 10 keywords
    return keywords[:10]


def classify_question_type(question: str) -> str:
    """Classify the type of question for better retrieval."""
    q_lower = question.lower()
    
    # How-to questions
    if any(phrase in q_lower for phrase in ['how to', 'how can', 'how do', 'how should']):
        return 'how-to'
    
    # Definition questions
    if any(phrase in q_lower for phrase in ['what is', 'what are', 'what does', 'define']):
        return 'definition'
    
    # Temporal questions
    if any(word in q_lower for word in ['when', 'what time', 'how long', 'duration']):
        return 'temporal'
    
    # Location questions
    if any(word in q_lower for word in ['where', 'which location']):
        return 'location'
    
    # Explanation questions
    if any(word in q_lower for word in ['why', 'reason', 'explain']):
        return 'explanation'
    
    # Permission/capability questions
    if any(phrase in q_lower for phrase in ['can i', 'is it possible', 'am i able', 'may i']):
        return 'permission'
    
    # Comparison questions
    if any(word in q_lower for word in ['or', 'versus', 'vs', 'difference between']):
        return 'comparison'
    
    # Yes/No questions
    if any(q_lower.startswith(word) for word in ['is ', 'are ', 'do ', 'does ', 'will ', 'can ']):
        return 'yes-no'
    
    # Default
    return 'general'

def load_and_split_txt(doc_path: str, chunk_size: int = 1000, chunk_overlap: int = 200):
    """Load TXT file and split into chunks."""
    RecursiveCharacterTextSplitter = _import_split_text()
    
    # Read text file
    with open(doc_path, 'r', encoding='utf-8') as f:
        text = f.read()
    
    # Split into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    chunks = text_splitter.create_documents([text])
    print(f"✓ Created {len(chunks)} chunks from: {doc_path}")
    return chunks


# =============================================================================
# Chunking plain text 
# =============================================================================

def split_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200):
    """split text into chunks."""
    RecursiveCharacterTextSplitter = _import_split_text()
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    chunks = text_splitter.split_text(text)
    print(f"✓ Created {len(chunks)} chunks from the text provided")
    return chunks


# =============================================================================
# Embedding & Vectorization
# =============================================================================

def create_vectorizer():
    """Create HuggingFace text vectorizer with embedding cache."""
    HFTextVectorizer, EmbeddingsCache = _import_vectorizer()
    
    # Initialize the cache first
    embed_cache = EmbeddingsCache(
        name="embedcache",
        redis_url=REDIS_URL,
        ttl=600
    )
    
    hf = HFTextVectorizer(
        # model="ai4bharat/indic-bert",
        model="sentence-transformers/all-MiniLM-L6-v2",
        cache=embed_cache
    )
    print("✓ Vectorizer ready")
    return hf


def embed_chunks(vectorizer, chunks) -> list:
    """Embed all document chunks."""
    print("Embedding chunks...")
    # Handle both string chunks and Document objects
    texts = []
    for chunk in chunks:
        if isinstance(chunk, str):
            texts.append(chunk)
        else:
            # Assume it's a Document object with page_content attribute
            texts.append(chunk.page_content)
    embeddings = vectorizer.embed_many(texts)
    assert len(embeddings) == len(chunks), "Embedding count mismatch"
    print(f"✓ Embedded {len(embeddings)} chunks")
    return embeddings


# =============================================================================
# Index Management
# =============================================================================

def create_async_index(schema: dict, llm_cache=None):
    
    # Surgical clears prevent "stale" cache hits
    if llm_cache: 
        llm_cache.clear()
    
    
    """Create Redis search index from schema."""
    AsyncSearchIndex, _, _, _ = _import_index()
    index_name = schema["index"]["name"]
    index = AsyncSearchIndex.from_dict(schema, redis_url=REDIS_URL)
    
    from redis import Redis
    r = Redis.from_url(REDIS_URL)
    # 1️⃣ Check if index already exists
    try:
        existing_indexes = r.execute_command("FT._LIST")
        existing_indexes = [i.decode() for i in existing_indexes]
        existing_indexes = [i.decode() for i in existing_indexes]

        if index_name in existing_indexes:
            print(f"✓ Index '{index_name}' already exists — returning existing")
            return index   # Just return the object without creating
    except Exception as e:
        print(f"Index existence check failed, attempting creation: {e}")
    
    index.create()
    print(f"✓ Async Index '{index_name}' created")
    return index


def load_data_to_index(index, chunks, embeddings) -> list:
    """Load embedded chunks into the Redis index."""
    _, _, _, array_to_buffer = _import_index()
    
    data = []
    for i, chunk in enumerate(chunks):
        # Handle both string chunks and Document objects
        if isinstance(chunk, str):
            content = chunk
        else:
            # Assume it's a Document object with page_content attribute
            content = chunk.page_content
        unique_id = str(uuid.uuid4())
        data.append({
            'chunk_id': unique_id,
            'content': content,
            'text_embedding': array_to_buffer(embeddings[i], dtype='float32')
        })
    
    keys = index.load(data, id_field="chunk_id")
    print(f"✓ Loaded {len(keys)} chunks into index")
    return keys

def load_data_to_index_new(index, chunks, embeddings) -> list:
    """Load embedded chunks with metadata into the Redis index."""
    _, _, _, array_to_buffer = _import_index()
    import hashlib
    
    data = []
    skipped = 0
    for i, chunk in enumerate(chunks):
        # Handle both string chunks and Document objects
        if isinstance(chunk, str):
            content = chunk
            metadata = {}  # No metadata for plain strings
        else:
            # Document object with page_content and metadata
            content = chunk.page_content
            metadata = chunk.metadata if hasattr(chunk, 'metadata') else {}
        
        # unique_id = str(uuid.uuid4())
        # Generate content hash for deduplication
        content_hash = hashlib.md5(content.encode('utf-8')).hexdigest()
        # Check if this content already exists
        try:
            existing = index.search(f"@chunk_id:{{{content_hash}}}")
            if existing.total > 0:
                skipped += 1
                print(f"  ⊘ Skipped duplicate chunk (hash: {content_hash[:8]}...)")
                continue
        except Exception as e:
            # Index might not exist yet or other error - continue loading
            pass
        # Base document
        doc = {
            # 'chunk_id': unique_id,
            'chunk_id': content_hash,  # Use hash instead of UUID
            'content': content,
            'text_embedding': array_to_buffer(embeddings[i], dtype='float32')
        }
        
        # Add metadata fields if available
        if metadata:
            doc['topic'] = metadata.get('topic', 'General')
            doc['keywords'] = ','.join(metadata.get('keywords', []))  # Convert list to comma-separated string
            doc['question_type'] = metadata.get('question_type', 'general')
            doc['question'] = metadata.get('question', '')
            doc['source_file'] = metadata.get('source_file', '')
        else:
            # Default values if no metadata
            doc['topic'] = 'General'
            doc['keywords'] = ''
            doc['question_type'] = 'general'
            doc['question'] = ''
            doc['source_file'] = ''
        
        data.append(doc)
    
    # keys = index.load(data, id_field="chunk_id")
    # print(f"✓ Loaded {len(keys)} chunks into index with metadata")
    if data:
        keys = index.load(data, id_field="chunk_id")
        print(f"✓ Loaded {len(keys)} new chunks into index")
    else:
        keys = []
        print(f"✓ No new chunks to load")
    
    if skipped > 0:
        print(f"⊘ Skipped {skipped} duplicate chunks")
    return keys

# =============================================================================
# Query Functions
# =============================================================================

def embed_query(query: str) -> list:
    """Convert a user query into a dense vector representation."""
    return _vectorizer.embed(query)

def user_query_caching(hf):
    from redisvl.extensions.cache.llm import SemanticCache

    llmcache = SemanticCache(
        name="llmcache",                
        vectorizer=hf,                  
        redis_url=REDIS_URL,          
        ttl=600,                         
        distance_threshold=0.12,        
        overwrite=True      
    )

    print("Cache created successfully:", llmcache)
    return llmcache

async def retrieve_context(
    async_index, 
    query_vector, 
    retrieval_distance_threshold=0.4,
    num_results=15
    ) -> str:
    """Fetch the relevant context from Redis using vector search."""
    _, _, VectorQuery, _ = _import_index()
    
    results = async_index.query(
        VectorQuery(
            vector=query_vector,
            vector_field_name="text_embedding",
            return_fields=["content", "vector_distance"],
            num_results=num_results
        )
    )
    print(results)

    formatted_context_list = []
    for result in results:
        distance_raw  = result.get("vector_distance")
        content = result.get("content", "")
        
        if distance_raw is not None : 
            distance = float(distance_raw)
            if distance <= retrieval_distance_threshold:
                formatted_context_list.append(content)
                print(f"  ✓ Chunk (distance={distance:.3f})")
            else :
                print(f"  ✗ Skipped (distance={distance:.3f} > {retrieval_distance_threshold})")
        else:
            print(f"  ✗ Skipped (distance={distance_raw:.3f} > {retrieval_distance_threshold})")

    return "\n\n".join(formatted_context_list)

from typing import Optional
async def retrieve_context_new(
    async_index, 
    query_vector, 
    query_text: str = "",  # NEW: Original query text for metadata filtering
    retrieval_distance_threshold: float = 0.4,
    num_results: int = 15,
    filter_by_topic: Optional[str] = None,  # NEW: Optional topic filter
    filter_by_question_type: Optional[str] = None,  # NEW: Optional question type filter
    use_metadata_boosting: bool = True,  # NEW: Boost results based on metadata match
) -> str:
    """
    Fetch relevant context from Redis using vector search with metadata-enhanced ranking.
    
    Args:
        async_index: Redis vector index
        query_vector: Query embedding vector
        query_text: Original query text (for metadata matching)
        retrieval_distance_threshold: Max vector distance to include results
        num_results: Number of results to retrieve
        filter_by_topic: Filter results by specific topic (e.g., "ACCOUNT OPENING")
        filter_by_question_type: Filter by question type (e.g., "how-to")
        use_metadata_boosting: Whether to boost scores based on metadata match
    
    Returns:
        Formatted context string with relevant FAQs
    """
    
    _, _, VectorQuery, _ = _import_index()
    
    # Updated return fields to include metadata
    return_fields = [
        "content", 
        "vector_distance",
        "topic",           # NEW
        "keywords",        # NEW
        "question_type",   # NEW
        "question",        # NEW
        "chunk_id"         # NEW
    ]
    
    results = async_index.query(
        VectorQuery(
            vector=query_vector,
            vector_field_name="text_embedding",
            return_fields=return_fields,
            num_results=num_results
        )
    )
    
    print(f"\n{'='*60}")
    print(f"RETRIEVAL RESULTS for query: '{query_text}'")
    print(f"{'='*60}")
    
    # Process and rank results with metadata
    enhanced_results = []
    
    for idx, result in enumerate(results, 1):
        distance_raw = result.get("vector_distance")
        content = result.get("content", "")
        topic = result.get("topic", "")
        keywords = result.get("keywords", "")
        question_type = result.get("question_type", "")
        question = result.get("question", "")
        chunk_id = result.get("chunk_id", "")
        
        # Validate distance
        try:
            if distance_raw is None:
                print(f"  ⚠ Chunk #{idx}: No distance found, skipping")
                continue
            distance = float(distance_raw)
        except (ValueError, TypeError) as e:
            print(f"  ⚠ Chunk #{idx}: Invalid distance '{distance_raw}', skipping")
            continue
        
        # Apply topic filter if specified
        if filter_by_topic and topic != filter_by_topic:
            print(f"  ✗ Chunk #{idx}: Topic mismatch ('{topic}' != '{filter_by_topic}')")
            continue
        
        # Apply question type filter if specified
        if filter_by_question_type and question_type != filter_by_question_type:
            print(f"  ✗ Chunk #{idx}: Question type mismatch ('{question_type}' != '{filter_by_question_type}')")
            continue
        
        # Calculate metadata boost score
        metadata_score = 0.0
        if use_metadata_boosting and query_text:
            metadata_score = calculate_metadata_boost(
                query_text, 
                topic, 
                keywords, 
                question_type,
                question
            )
        
        # Adjusted score: lower is better (distance - boost)
        adjusted_distance = distance - metadata_score
        
        enhanced_results.append({
            "content": content,
            "distance": distance,
            "adjusted_distance": adjusted_distance,
            "topic": topic,
            "keywords": keywords,
            "question_type": question_type,
            "question": question,
            "chunk_id": chunk_id,
            "metadata_score": metadata_score
        })
        
        print(f"  • Chunk #{idx} [{chunk_id}]")
        print(f"    Distance: {distance:.3f} | Adjusted: {adjusted_distance:.3f} | Boost: +{metadata_score:.3f}")
        print(f"    Topic: {topic} | Type: {question_type}")
        print(f"    Question: {question[:80]}...")
    
    # Sort by adjusted distance (lower is better)
    enhanced_results.sort(key=lambda x: x["adjusted_distance"])
    
    # Filter by threshold and format
    formatted_context_list = []
    print(f"\n{'='*60}")
    print(f"FILTERED RESULTS (threshold={retrieval_distance_threshold})")
    print(f"{'='*60}")
    
    for idx, result in enumerate(enhanced_results, 1):
        if result["adjusted_distance"] <= retrieval_distance_threshold:
            formatted_context_list.append(result["content"])
            print(f"  ✓ [{result['chunk_id']}] {result['question'][:60]}...")
            print(f"    Distance: {result['distance']:.3f} → {result['adjusted_distance']:.3f} | Topic: {result['topic']}")
        else:
            print(f"  ✗ [{result['chunk_id']}] Skipped (distance={result['adjusted_distance']:.3f} > {retrieval_distance_threshold})")
    
    print(f"\n✓ Selected {len(formatted_context_list)} / {len(enhanced_results)} results")
    print(f"{'='*60}\n")
    
    return "\n\n".join(formatted_context_list)


def calculate_metadata_boost(
    query: str, 
    topic: str, 
    keywords: str, 
    question_type: str,
    question: str
) -> float:
    """
    Calculate a boost score based on metadata matching.
    Higher score = better match = lower adjusted distance.
    
    Returns: Float between 0.0 and 0.15 (significant but not overwhelming)
    """
    boost = 0.0
    query_lower = query.lower()
    
    # Keywords matching (most important) - up to 0.08 boost
    if keywords:
        keyword_list = [k.lower() for k in keywords.split(',') if k]
        matching_keywords = sum(1 for kw in keyword_list if kw in query_lower)
        if matching_keywords > 0:
            boost += min(matching_keywords * 0.02, 0.08)
    
    # Topic matching - 0.04 boost
    if topic and topic.lower() in query_lower:
        boost += 0.04
    
    # Question type matching - 0.03 boost
    if question_type:
        type_indicators = {
            'how-to': ['how to', 'how can', 'how do'],
            'definition': ['what is', 'what are', 'define'],
            'temporal': ['when', 'how long'],
            'explanation': ['why', 'reason'],
            'permission': ['can i', 'is it possible'],
        }
        
        if question_type in type_indicators:
            if any(indicator in query_lower for indicator in type_indicators[question_type]):
                boost += 0.03
                
    # NEW: Direct question similarity boost - up to 0.05
    if question:
        question_lower = question.lower()
        # Check if query words appear in the stored question
        query_words = set(query_lower.split())
        question_words = set(question_lower.split())
        common_words = query_words & question_words
        
        if len(common_words) > 2:  # At least 3 words match
            boost += 0.05
    
    return min(boost, 0.15)  # Cap at 0.15 to avoid over-boosting

# =============================================================================
# LLM Response Generation
# =============================================================================

def promptify(query: str, context: str, domain: str = "general") -> str:
    # Generates a language-agnostic prompt for cross-lingual RAG
    return f'''You are an expert {domain} assistant. 
    Review the context blocks below to answer the question.
    
    STRICT INSTRUCTION:
    - Keep the response short crisp to the point. 
    - Respond in the language of the User Question ; default as english . 
    - Every factual claim MUST be followed by its original phrase from the context in brackets [Original: "..."].
    - Base your answer ONLY on the provided context.
    - If you cannot answer based on the context, do not guess.
    
    User question:
    {query}

    Helpful context:
    {context}

    Answer:
    '''




async def refine_with_slm(user_question, context_chunks):
    import requests
    # Default to localhost for local testing
    OLLAMA_URL = "http://localhost:11434"
    """
    Takes a list of RAG strings and uses local Qwen to refine them.
    """
    print("--- Sending to Local SLM ---")
    
    
    # 1. Format the chunks into a single text block
    # formatted_context = "\n\n".join([f"Source {i+1}: {chunk}" for i, chunk in enumerate(context_chunks)])
    final_prompt = promptify(user_question, context_chunks) # Example of promptify logic
    print(f"final_prompt : {final_prompt}")
    
    # 2. Construct the Payload
    payload = {
        "model": "qwen2.5:1.5b",  # Must match the model you pulled
        "messages": [
            {
                "role": "system", 
                "content": "You are a helpful assistant. Use ONLY the provided Sources to answer. If the answer is not in the Sources, say 'I don't know'."
            },
            {
                "role": "user", 
                "content": f"{final_prompt}"
            }
        ],
        "stream": False,
        "temperature": 0.1 # Low temp = more factual
    }

    # 3. Send to Ollama
    try:
        response = requests.post(f"{OLLAMA_URL}/api/chat", json=payload)
        
        if response.status_code == 200:
            result = response.json()
            answer = result['message']['content']
            print(f"--- SLM Answer: {answer} ---")
            return answer
        else:
            print(f"Error: {response.text}")
            return "Sorry, I had trouble processing that."
            
    except Exception as e:
        print(f"Connection Error: {e}")
        return "System is currently offline."

async def generate_llm_response(query: str, context: str) -> str:
    """Send prompt to Anthropic LLM and return response."""
    from google import genai
    from google.genai import types
    gemini = genai
    client = gemini.Client(api_key=gemini_api_key).aio # Get the asynchronous client (client.aio)

    # 2. Prepare the prompt 
    final_prompt = promptify(query, context) # Example of promptify logic

    # 3. Define the configuration for the call
    generation_config = types.GenerateContentConfig(
        # The system instruction/prompt goes here
        system_instruction=SYSTEM_PROMPT,
        max_output_tokens=2048,
        temperature=0.15,
    )
    
    CHAT_MODEL = "gemini-2.5-flash"
    # 4. Make the asynchronous API call
    response = await client.models.generate_content(
        model=CHAT_MODEL,
        contents=[final_prompt], # Contents is the list of user prompts/parts
        config=generation_config
    )   
    
    # 5. Extract the text response
    # The Gemini response object has a simple .text attribute
    return response.text



async def answer_question_old(index, query: str , cache, use_cache=True, use_llm=True) -> str:
    """End-to-end RAG: embeds query, retrieves context, generates LLM response."""
    query_vector = embed_query(query)
    results = []
    
    if use_cache :
        results = cache.check(vector=query_vector)
    
    if results:
        print("found similar, semantic")
        return results[0]['response']

    context = await retrieve_context_new(index, query_vector, query)

    if not use_llm : 
        return context
    llmResults = await generate_llm_response(query, context)
    # llmResults = await refine_with_slm(query, context)
    cache.store(query, llmResults, query_vector)
    return llmResults


def translate_py(text : str , source_lang : str = "en" , target_lang : str = "en") :
    from translatepy import Translator
    try : 
        # Initialize translator
        translator = Translator()
        print(f"🔄 Translating from {source_lang} to {target_lang}")
        result = translator.translate(text, target_lang, source_lang)
        final_text = result.result
        print(f"✅ Translated: {final_text}")
        return final_text
    except Exception as trans_error:
        print(f"⚠️ Translation failed, using original text: {trans_error}")
        final_text = text
        return final_text
    
from deep_translator import GoogleTranslator
# Initialize outside to reuse the connection session
def translate_google(text: str, source_lang: str = "auto", target_lang: str = "en"):
    print(f"🔄 Translating {text[:20]}... from {source_lang} to {target_lang}")
    try:
        # 'auto' is better for the source if you aren't 100% sure it's English
        translator = GoogleTranslator(source=source_lang, target=target_lang)
        final_text = translator.translate(text)
        
        if not final_text:
            return text
            
        print(f"✅ Translated: {final_text}")
        return final_text
    except Exception as trans_error:
        print(f"⚠️ Translation failed: {trans_error}")
        return text
    
async def answer_question(
    index, 
    query: str,
    cache, 
    use_cache=True, 
    use_llm=True,
    original_query: str = None,
    source_lang: str = "en",
    translate_rag: bool = False  # NEW
):
    """End-to-end RAG: embeds query, retrieves context, generates LLM response."""
    query_vector = embed_query(query)
    results = []
    
    if use_cache:
        results = cache.check(vector=query_vector)
    
    if results:
        print("✅ Found in cache")
        cached_response = results[0]['response']
        
        # If cached and needs translation
        if source_lang != "en" and (use_llm or translate_rag):
            if use_llm and original_query:
                # Translate cached response to source language
                translated = translate_google(cached_response, "en", source_lang)
                if translate_rag:
                    return {
                        'answer': translated,
                        'original_rag_context': cached_response,
                        'translated_rag_context': translated
                    }
                return translated
            elif translate_rag:
                translated = translate_google(cached_response, "en", source_lang)
                return {
                    'answer': translated,
                    'original_rag_context': cached_response,
                    'translated_rag_context': translated
                }
        return cached_response

    # Retrieve context from RAG
    context = await retrieve_context_new(index, query_vector, query)

    # If NOT using LLM, return RAG context (optionally translated)
    if not use_llm:
        if source_lang != "en" and translate_rag:
            translated_context = translate_google(context, "en", source_lang)
            return {
                'answer': translated_context,
                'original_rag_context': context,
                'translated_rag_context': translated_context
            }
        # elif source_lang != "en":
        #     # Translate only the answer
        #     return translate_google(context, "en", source_lang)
        return context

    if original_query : 
        llmResults = await generate_llm_response(original_query, context)
        cache.store(original_query, llmResults, embed_query(query))
    else : 
        llmResults = await generate_llm_response(query, context)
        cache.store(query, llmResults, query_vector)
    
    # Return with RAG context if requested
    if translate_rag and source_lang != "en":
        translated_context = translate_google(context, "en", source_lang)
        return {
            'answer': llmResults,
            'original_rag_context': context,
            'translated_rag_context': translated_context
        }
    
    return llmResults
# =============================================================================
# Initialization
# =============================================================================

def initialize(doc_path: str = DOC_PATH):
    """Initialize the RAG system."""
    global _vectorizer, _async_index
    _, AsyncSearchIndex, _, _ = _import_index()

    print("\n" + "=" * 50)
    print("Initializing RAG System")
    print("=" * 50)

    # 1. Test Redis
    get_redis_connection()

    # 2. Load document
    #chunks = load_and_split_docx_new(doc_path)
    chunks = load_and_split_document(doc_path)

    # 3. Create vectorizer
    _vectorizer = create_vectorizer()
    
    llmCache = user_query_caching(_vectorizer)

    # 4. Embed chunks
    embeddings = embed_chunks(_vectorizer, chunks)
    
    # 5. Create & load index (Updated to clear caches surgically)
    # Extract the embedding cache from the vectorizer object
    embed_cache = _vectorizer.cache if hasattr(_vectorizer, 'cache') else None

    # 5. Create & load index
    # async_index = create_async_index(SCHEMA)
    async_index = create_async_index(
        SCHEMA, 
        embed_cache=embed_cache, 
        llm_cache=llmCache
    )
    load_data_to_index_new(async_index, chunks, embeddings)
    #load_data_to_index()

    print("=" * 50)
    print("✓ RAG System Ready!")
    print("=" * 50 + "\n")

    return async_index,llmCache


def ensure_api_key():
    """Ensure Anthropic API key is set."""
    if "ANTHROPIC_API_KEY" not in os.environ:
        os.environ["ANTHROPIC_API_KEY"] = getpass.getpass("ANTHROPIC_API_KEY: ")
    print(f"Using model: {CHAT_MODEL}")


# async def run_test_questions(async_index, questions: list , cache):
#     """Run test questions through the RAG system."""
#     results = await asyncio.gather(*[
#         answer_question(async_index, q , cache) for q in questions
#     ])
#     #questions.append("बोनस शेयर्स कैसे और कहां दिखते हैं", "आईपीओ के लिए आवेदन कैसे करें", "मुझे लॉगिन क्रेडेंशियल कब मिलेंगे?")
#     for i, result in enumerate(results):
#         print(f"\nQ: {questions[i]}")
#         print(f"A: {result}")
#         print("-" * 50)
async def run_test_questions(async_index, test_cases: list, cache):
    """Run test questions through the RAG system with varying parameters."""
    
    tasks = []
    for tc in test_cases:
        if isinstance(tc, str):
            # Simple case: only the query string is provided
            tasks.append(answer_question(async_index, tc, cache))
        elif isinstance(tc, dict):
            # Complex case: Dictionary contains 'query' and other flags
            # We pop 'query' to pass it positionally, then unpack the rest
            query = tc.pop('query')
            tasks.append(answer_question(async_index, query, cache, **tc))

    results = await asyncio.gather(*tasks)

    for i, result in enumerate(results):
        # Determine what to print as the question header
        current_case = test_cases[i]
        q_text = current_case if isinstance(current_case, str) else current_case.get('original_query', "Query")
        
        print(f"\nQ: {q_text}")
        print(f"A: {result}")
        print("-" * 50)


# =============================================================================
# Main
# =============================================================================

def main():
    """Main entry point."""
    print("RAG PoC Starting...")
    
    # Ensure API key
    #ensure_api_key()

    # Initialize (lazy loading happens here)
    #async_index,llmCache = initialize("resources/sharekhan_faqs_complete.docx")
    async_index,llmCache = initialize("C:/Users/tanishka.kothari/Downloads/1769772850270.pdf")
    
    # Test
    # questions = [
    #     "Why are transferred funds not credited In the account?",
    #     {
    #         "query": translate_google("आईपीओ के लिए आवेदन कैसे करें", "hi"),
    #         "original_query": "आईपीओ के लिए आवेदन कैसे करें",
    #         "source_lang": "hi",
    #         "translate_rag": True
    #     },
    #     {
    #         "query": translate_google("मुझे लॉगिन क्रेडेंशियल कब मिलेंगे?", "hi"),
    #         "original_query": "मुझे लॉगिन क्रेडेंशियल कब मिलेंगे?",
    #         "source_lang": "hi"
    #     },{
    #         "query": translate_google("बोनस शेयर्स कैसे और कहां दिखते हैं", "hi"),
    #         "original_query": "बोनस शेयर्स कैसे और कहां दिखते हैं",
    #         "source_lang": "hi",
    #         "translate_rag": True,
    #         "use_llm": True
    #     }
    # ]
    
    questions = [
        "Who can use this special window?"
    ]

    asyncio.run(run_test_questions(async_index, questions, llmCache))




if __name__ == "__main__":
    main()